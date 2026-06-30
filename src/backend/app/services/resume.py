import io
import re
import httpx
import pdfplumber
from docx import Document
from typing import Optional, Tuple
import logging

logger = logging.getLogger(__name__)

def normalize_drive_link(url: str) -> Optional[str]:
    """
    Extracts the file ID from a Google Drive share link and constructs
    a direct download URL.
    """
    if not url or not isinstance(url, str):
        return None
    
    # Matches patterns like /file/d/<ID>/view or ?id=<ID> or /open?id=<ID>
    patterns = [
        r"/file/d/([a-zA-Z0-9_-]+)",
        r"id=([a-zA-Z0-9_-]+)",
        r"/open\?id=([a-zA-Z0-9_-]+)"
    ]
    
    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            file_id = match.group(1)
            return f"https://drive.google.com/uc?export=download&id={file_id}"
            
    return None

def download_resume(direct_url: str) -> bytes:
    """
    Downloads the resume content. Handles Google Drive virus scan warnings
    by extracting the confirmation code and retrying the request.
    """
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
    }
    
    with httpx.Client(follow_redirects=True, headers=headers, timeout=30.0) as client:
        response = client.get(direct_url)
        
        # Check if the page is a Google Drive virus scan warning
        # It usually contains a "confirm=" query string in a link or form
        if "confirm=" in response.text:
            match = re.search(r'confirm=([a-zA-Z0-9_-]+)', response.text)
            if match:
                confirm_code = match.group(1)
                # Retry download appending confirmation code
                confirm_url = f"{direct_url}&confirm={confirm_code}"
                response = client.get(confirm_url)
                
        if response.status_code != 200:
            raise ValueError(f"HTTP download failed with status code {response.status_code}")
            
        return response.content

def extract_text_from_pdf(content: bytes) -> str:
    """
    Extracts text from PDF bytes using pdfplumber.
    """
    text_pages = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_pages.append(page_text)
    return "\n".join(text_pages).strip()

def extract_text_from_docx(content: bytes) -> str:
    """
    Extracts text from DOCX bytes using python-docx.
    """
    doc = Document(io.BytesIO(content))
    text_runs = [para.text for para in doc.paragraphs]
    return "\n".join(text_runs).strip()

def process_resume(url: str) -> Tuple[str, Optional[str]]:
    """
    Processes candidate resume: normalizes Google Drive link, downloads it,
    detects format (PDF/DOCX), extracts text, and returns status and text.
    Returns:
        Tuple[str, Optional[str]]: (resume_status, extracted_text)
        Where resume_status is 'EXTRACTED' or 'UNAVAILABLE'
    """
    direct_url = normalize_drive_link(url)
    if not direct_url:
        logger.warning(f"Could not normalize Google Drive link: {url}")
        return "UNAVAILABLE", None

    try:
        content = download_resume(direct_url)
        
        # Determine format based on magic bytes or file headers
        # PDF starts with %PDF- (hex: 25 50 44 46)
        # DOCX starts with PK.. (hex: 50 4B 03 04) - standard zip archive
        if content.startswith(b"%PDF"):
            extracted_text = extract_text_from_pdf(content)
        elif content.startswith(b"PK\x03\x04"):
            extracted_text = extract_text_from_docx(content)
        else:
            # Fallback attempts
            try:
                extracted_text = extract_text_from_pdf(content)
            except Exception:
                extracted_text = extract_text_from_docx(content)

        if not extracted_text:
            # Try OCR fallback using pytesseract if available
            try:
                import pytesseract
                from PIL import Image
                import pypdfium2 as pdfium
                
                # Render PDF pages to images for OCR
                pdf = pdfium.PdfDocument(io.BytesIO(content))
                text_pages = []
                for i in range(len(pdf)):
                    page = pdf[i]
                    bitmap = page.render(scale=2)
                    pil_img = bitmap.to_pil()
                    page_text = pytesseract.image_to_string(pil_img)
                    if page_text:
                        text_pages.append(page_text)
                extracted_text = "\n".join(text_pages).strip()
            except Exception as ocr_err:
                logger.warning(f"OCR Fallback skipped or failed: {str(ocr_err)}")
                extracted_text = None

        if extracted_text and len(extracted_text) > 20:
            return "EXTRACTED", extracted_text
        else:
            logger.warning("Extracted resume text is empty or too short.")
            return "UNAVAILABLE", None

    except Exception as e:
        logger.error(f"Failed to process candidate resume from link {url}: {str(e)}")
        return "UNAVAILABLE", None
